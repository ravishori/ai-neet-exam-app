"""Convert Volume 1 Markdown master to a DOCX-ready Word document."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

SRC = Path(__file__).resolve().parent / "VOLUME_01_EXECUTIVE_PRODUCT_BLUEPRINT.md"
OUT = Path(__file__).resolve().parent / "TALOS-VOL-01-Executive-Product-Blueprint.docx"


def strip_yaml(text: str) -> str:
    if not text.lstrip().startswith("---"):
        return text
    rest = text.lstrip()[3:]
    end = rest.find("\n---")
    if end == -1:
        return text
    return rest[end + 4 :].lstrip("\n")


def add_runs(paragraph, text: str) -> None:
    # Very small markdown inline: **bold** and `code`
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def main() -> None:
    raw = strip_yaml(SRC.read_text(encoding="utf-8"))
    lines = raw.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_code = False
    code_lines: list[str] = []
    table_buf: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buf
        if not table_buf:
            return
        # filter separator rows
        rows = [r for r in table_buf if not all(re.fullmatch(r":?-+:?", c) for c in r)]
        table_buf = []
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j in range(cols):
                cell = table.rows[i].cells[j]
                cell.text = row[j] if j < len(row) else ""
                if i == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
        doc.add_paragraph()

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        code_lines = []

    for line in lines:
        if line.strip() == "\\newpage":
            flush_table()
            flush_code()
            doc.add_page_break()
            continue

        if line.startswith("```"):
            flush_table()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            table_buf.append(cells)
            continue
        else:
            flush_table()

        if not line.strip():
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_runs(p, line[2:].strip())
            if p.runs:
                p.runs[0].italic = True
        elif re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\. ", "", line).strip())
        else:
            p = doc.add_paragraph()
            add_runs(p, line.strip())

    flush_table()
    flush_code()

    # Cover alignment tweak for first heading if present
    if doc.paragraphs:
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
